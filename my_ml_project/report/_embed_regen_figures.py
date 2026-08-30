"""
Embeds the 3 regenerated PNGs (figure_4_1_workflow.png, figure_5_1_dataset.png,
figure_5_2_dataset.png) into the image paragraphs immediately preceding their
respective captions in both Main-Book.docx and Report-02-UPDATED.docx.

Replaces the image part's blob in place (same rId, same content-type — PNG to
PNG) and rescales the displayed extent to match the new image's true aspect
ratio, capped at a sensible max width/height matching the rest of the document.
"""
import os
import zipfile
from lxml import etree
from docx import Document
from PIL import Image

REPORT_DIR = r"d:\laragon\www\cse_fresh_graduate_predict_employable_ml_model\my_ml_project\report"

EMU_PER_INCH = 914400
MAX_WIDTH_IN = 6.2
MAX_HEIGHT_IN = 8.5

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}

REPLACEMENTS = {
    "Main-Book.docx": [
        (359, "figure_5_7_console.png"),
        (364, "figure_5_8_confusion_matrix.png"),
        (366, "figure_5_9_confusion_matrix.png"),
    ],
    "Report-02-UPDATED.docx": [
        (305, "figure_5_7_console.png"),
        (310, "figure_5_8_confusion_matrix.png"),
        (312, "figure_5_9_confusion_matrix.png"),
    ],
}


def get_extent_emu(new_path):
    with Image.open(new_path) as img:
        w_px, h_px = img.size
    aspect = h_px / w_px
    width_in = MAX_WIDTH_IN
    height_in = width_in * aspect
    if height_in > MAX_HEIGHT_IN:
        height_in = MAX_HEIGHT_IN
        width_in = height_in / aspect
    return int(width_in * EMU_PER_INCH), int(height_in * EMU_PER_INCH)


def replace_image_in_paragraph(doc, para_idx, new_image_path):
    p = doc.paragraphs[para_idx]
    p_xml = p._p
    blips = p_xml.findall(".//a:blip", NS)
    if not blips:
        raise RuntimeError(f"No blip found in paragraph {para_idx}")
    blip = blips[0]
    rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
    image_part = p.part.related_parts[rId]

    with open(new_image_path, "rb") as f:
        new_bytes = f.read()
    image_part._blob = new_bytes

    cx, cy = get_extent_emu(new_image_path)

    for ext in p_xml.findall(".//wp:extent", NS):
        ext.set("cx", str(cx))
        ext.set("cy", str(cy))
    for ext in p_xml.findall(".//a:ext", NS):
        ext.set("cx", str(cx))
        ext.set("cy", str(cy))

    return rId


def main():
    for fname, replacements in REPLACEMENTS.items():
        path = os.path.join(REPORT_DIR, fname)
        doc = Document(path)
        before_paras = len(doc.paragraphs)
        before_tables = len(doc.tables)

        for para_idx, image_name in replacements:
            image_path = os.path.join(REPORT_DIR, image_name)
            rId = replace_image_in_paragraph(doc, para_idx, image_path)
            print(f"{fname}: paragraph {para_idx} <- {image_name} (rId={rId})")

        doc.save(path)

        # Verify structural integrity
        doc2 = Document(path)
        after_paras = len(doc2.paragraphs)
        after_tables = len(doc2.tables)
        assert before_paras == after_paras, f"Paragraph count changed in {fname}: {before_paras} -> {after_paras}"
        assert before_tables == after_tables, f"Table count changed in {fname}: {before_tables} -> {after_tables}"

        with zipfile.ZipFile(path) as z:
            xml_bytes = z.read("word/document.xml")
            etree.fromstring(xml_bytes)
        print(f"{fname}: OK, paragraphs={after_paras}, tables={after_tables}, XML valid")


if __name__ == "__main__":
    main()
