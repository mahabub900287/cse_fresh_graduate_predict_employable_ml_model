"""
Phase 3: applies the Phase 2 audit's verified numbers to the thesis text —
adds a cross-validation/robustness subsection to Section 5.7, corrects two
stale limitation paragraphs in Section 6.2 that still described the
abandoned SMOTE/0.35-threshold configuration and an outdated "only XGBoost
was compared" claim, softens the "Resume_Score dominant" framing in Section
5.10 with the permutation-importance finding, and expands Section 5.14's
limitations list with the new CV/fairness/error-analysis findings.

All inserted numbers are taken directly from
d:\\laragon\\www\\cse_fresh_graduate_predict_employable_ml_model\\my_ml_project\\phase2_*.json
(the actual outputs of the Phase 2 audit scripts), not invented.

Run against both Main-Book.docx and Report-02-UPDATED.docx, whose bodies are
otherwise word-for-word identical for the sections touched here.
"""
import copy
import docx

PATHS = [
    r"d:\laragon\www\cse_fresh_graduate_predict_employable_ml_model\my_ml_project\book\Main-Book.docx",
    r"d:\laragon\www\cse_fresh_graduate_predict_employable_ml_model\my_ml_project\book\Report-02-UPDATED.docx",
]


def full_text(p):
    return "".join(r.text for r in p.runs)


def replace_paragraph_text(paragraph, old, new, required=True):
    full = full_text(paragraph)
    if old not in full:
        if required:
            raise AssertionError(f"OLD text not found: {old[:80]!r}")
        return False
    newfull = full.replace(old, new)
    paragraph.runs[0].text = newfull
    for r in paragraph.runs[1:]:
        r.text = ""
    return True


def find_paragraph_index(doc, predicate):
    for i, p in enumerate(doc.paragraphs):
        if predicate(full_text(p).strip()):
            return i
    return None


def insert_paragraph_after(reference_paragraph, text="", style=None, bold_first_sentence=None):
    new_element = copy.deepcopy(reference_paragraph._p)
    for child in list(new_element):
        if child.tag.endswith("}r"):
            new_element.remove(child)
    reference_paragraph._p.addnext(new_element)
    from docx.text.paragraph import Paragraph
    new_paragraph = Paragraph(new_element, reference_paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    if text:
        run = new_paragraph.add_run(text)
    return new_paragraph


# ---------------------------------------------------------------------------
# New Section 5.7 subsection: Cross-Validation and Robustness Validation
# (inserted after the ROC-AUC paragraph, before Section 5.8 heading)
# ---------------------------------------------------------------------------
CV_HEADING_TEXT = "5.7.1 Cross-Validation and Robustness Validation"

CV_PARA_1 = (
    "The single stratified 80:20 hold-out reported above was additionally stress-tested against three "
    "independent robustness checks, none of which altered the reported configuration or headline figures, "
    "so as to establish that the result is a stable population-level estimate rather than an artefact of one "
    "particular split. First, stratified 5-fold cross-validation was run on the full 50,000-record dataset for "
    "the tuned XGBoost pipeline: mean accuracy 0.8176 (standard deviation 0.0027) and mean ROC-AUC 0.8075 "
    "(standard deviation 0.0062) across the five folds, both figures comfortably encompassing the single-split "
    "test result of 0.8170 accuracy and 0.8068 AUC reported in Table 5.5. The same 5-fold procedure repeated on "
    "the training partition alone \u2014 the population from which the RandomizedSearchCV's internal 3-fold search "
    "actually draws \u2014 gave 0.8182 (\u00b10.0040) accuracy and 0.8083 (\u00b10.0033) AUC, again consistent. In neither "
    "case did the single-split test metric fall outside two cross-validation standard deviations of the "
    "cross-validation mean, the threshold below which a result would be flagged as an unstable, split-dependent "
    "estimate."
)

CV_PARA_2 = (
    "Second, the tuned pipeline was re-fitted with its existing hyperparameters at four additional random seeds "
    "(7, 123 and 2024, alongside the reported seed 42) for the train/test split only, with no re-tuning. Test "
    "accuracy ranged from 0.8163 to 0.8191 and test ROC-AUC from 0.8041 to 0.8079 across the four seeds \u2014 a "
    "spread of 0.28 and 0.38 percentage points respectively \u2014 confirming the reported result is not sensitive "
    "to the particular seed value chosen. Third, the RandomizedSearchCV hyperparameter search itself was re-run "
    "with a substantially larger budget (40 candidate combinations against the original 15, same search space "
    "and 3-fold cross-validation), evaluated once on the untouched test set after the wider search concluded: "
    "test accuracy moved by \u22120.01 percentage points and test ROC-AUC by +0.06 percentage points relative to the "
    "original 15-candidate search, indicating the original search budget was already sufficient and the reported "
    "configuration is not meaningfully under-tuned."
)

CV_PARA_3 = (
    "Finally, the gap between training-set and held-out test-set performance was examined directly as a check "
    "for overfitting: the tuned XGBoost pipeline scored 0.8219 training accuracy against 0.8170 test accuracy, a "
    "gap of 0.5 percentage points, and 0.8172 training ROC-AUC against 0.8068 test ROC-AUC, a gap of 1.0 "
    "percentage points \u2014 both consistent with the cross-validation mean and well short of the gap that would "
    "indicate memorisation of the training data. This is not true of every algorithm compared in Table 5.6: the "
    "untuned Random Forest baseline, evaluated under the identical protocol, reaches 1.0000 training accuracy "
    "and ROC-AUC (a train\u2013test gap of 18.6 and 20.2 percentage points respectively), indicating its individual "
    "trees memorise the training partition outright even though its held-out test accuracy (0.8138) remains "
    "unremarkable and close to the other models; this caveat is recorded in Table 5.6 and should be read "
    "alongside that table's headline comparison. Logistic Regression and Decision Tree show train\u2013test gaps of "
    "0.1 and 0.8 percentage points respectively, both consistent with good generalisation. Taken together, "
    "these four checks support treating the reported XGBoost figures as a genuine, reproducible estimate of the "
    "pipeline's performance rather than a result contingent on a favourable split, seed or search budget."
)


# ---------------------------------------------------------------------------
# Section 5.10: soften the "Resume_Score alone accounts for 33.3 per cent"
# framing with the permutation-importance finding (new paragraph inserted
# after the existing Overall_Preparedness_Index caveat paragraph)
# ---------------------------------------------------------------------------
PERM_IMPORTANCE_PARA = (
    "This gain-based ranking was additionally cross-checked against permutation importance \u2014 a model-agnostic "
    "measure that shuffles one feature at a time in the held-out test set and records the resulting drop in "
    "ROC-AUC, so that a feature's importance is measured by how much the model's held-out performance actually "
    "depends on it, rather than by how often or how early it is used to split trees during training. The two "
    "measures agree closely on which features matter at all: the same four attributes \u2014 Resume_Score, "
    "Overall_Preparedness_Index, Programming_Skill and Internships \u2014 occupy the top four positions under both "
    "methods, and the bottom eight features (GitHub_Profile, Leadership_Experience, Academic_Performance, "
    "English_Proficiency, Teamwork, Skills_per_Project, Hackathons and Study_Hours_Per_Week) show permutation "
    "importances indistinguishable from zero under both methods. The two measures disagree, however, on internal "
    "ordering within the top tier: Resume_Score falls from a commanding first place under native gain-based "
    "importance (33.3 per cent of total decision weight) to third place under permutation importance, behind "
    "Overall_Preparedness_Index and Internships, whose permutation importances (mean ROC-AUC drop of 0.0134 and "
    "0.0081 respectively, against 0.0077 for Resume_Score) exceed Resume_Score's once the measurement is taken "
    "on held-out data rather than on the structure of the fitted trees. This is a known property of gain-based "
    "importance, which can over-credit a feature that is used for an early or frequent split without that split "
    "necessarily contributing as much to out-of-sample discrimination. Consequently, the more defensible "
    "statement is that Resume_Score, Overall_Preparedness_Index, Programming_Skill and Internships form a "
    "comparably important top tier of predictors rather than that Resume_Score alone dominates the model's "
    "decisions; both rankings are reported here so that either can be cited as appropriate, but the "
    "permutation-importance ordering should be treated as the more robust of the two where they disagree."
)


# ---------------------------------------------------------------------------
# Section 5.14: corrected limitation count and two new limitations
# (Major-linked error disparity; false-negative resume/interview pattern)
# ---------------------------------------------------------------------------
OLD_LIM_INTRO = "Six limitations emerge specifically from the results rather than from the study design in the abstract."
NEW_LIM_INTRO = "Eight limitations emerge specifically from the results rather than from the study design in the abstract."

OLD_LIM_FOURTH = (
    "Fourth, the comparative evaluation underlying RQ3 was run only against Logistic Regression, Decision Tree "
    "and Random Forest; Support Vector Machine was excluded on computational-cost grounds, so the claim that no "
    "algorithm dominates is established among four of the five originally identified candidates."
)
NEW_LIM_FOURTH = (
    "Fourth, the comparative evaluation underlying RQ3 was run against Logistic Regression, Decision Tree, "
    "Random Forest and XGBoost; Support Vector Machine alone was excluded on computational-cost grounds, so the "
    "claim that no algorithm dominates is established among four of the five originally identified candidates."
)

OLD_LIM_TAIL = (
    "Sixth, the cross-sectional nature of the data means the ranked determinants of Section 5.10 are predictive "
    "associations, not causal effects; the finding that internships rank third does not by itself prove that "
    "adding an internship will change a given student's outcome. Each of these limitations directly motivates a "
    "corresponding item in the future work of Chapter 6."
)
NEW_LIM_TAIL = (
    "Sixth, the cross-sectional nature of the data means the ranked determinants of Section 5.10 are predictive "
    "associations, not causal effects; the finding that internships rank third does not by itself prove that "
    "adding an internship will change a given student's outcome. Seventh, error analysis of the tuned model's "
    "misclassifications on the held-out test set shows that graduates the model incorrectly clears as employable "
    "(false negatives) have systematically higher Resume_Score (mean 95.8 versus 79.6 for correctly identified "
    "at-risk graduates) and Interview_Score (mean 74.3 versus 56.0) than the at-risk graduates it correctly "
    "flags, indicating the model is specifically prone to being misled by a strong-looking resume and interview "
    "record on students who are nonetheless not placed \u2014 a pattern consistent with the ranking caveat above and "
    "one that should inform how much weight any deployment places on Resume_Score in isolation. Eighth, although "
    "Major, Gender and University_Year are excluded from training, an audit-only cross-tabulation of prediction "
    "errors against these attributes shows accuracy varies materially by Major \u2014 from 71.3 per cent (Business "
    "Analytics) and 71.9 per cent (Electrical Engineering) to 86.0\u201386.8 per cent (Artificial Intelligence, "
    "Software Engineering) \u2014 while Gender (81.4\u201382.1 per cent) and University_Year (80.4\u201383.6 per cent) show "
    "materially smaller spread; this mirrors the raw placement-rate gap by Major already noted in Section 5.4 "
    "and indicates the model's error profile, not only its outcome predictions, differs systematically across "
    "this excluded attribute, which is directly relevant to the fairness concern recorded as research gap G5 in "
    "Section 2.7 and is not yet mitigated by excluding Major from training alone. Each of these limitations "
    "directly motivates a corresponding item in the future work of Chapter 6."
)


def insert_toc_entry_after(doc, after_predicate, new_text):
    """Clone the ToC line matching after_predicate and set its text to new_text,
    inserting it immediately after the matched line."""
    idx = find_paragraph_index(doc, after_predicate)
    assert idx is not None, "ToC anchor line not found"
    ref_p = doc.paragraphs[idx]
    new_el = copy.deepcopy(ref_p._p)
    ref_p._p.addnext(new_el)
    from docx.text.paragraph import Paragraph
    new_p = Paragraph(new_el, ref_p._parent)
    if new_p.runs:
        new_p.runs[0].text = new_text
        for r in new_p.runs[1:]:
            r.text = ""
    return idx


def process_file(path):
    doc = docx.Document(path)

    # 0. Add a ToC entry for the new 5.7.1 subsection, right after the
    # "5.7 Model Performance and Evaluation" front-matter ToC line and before
    # "5.8 Justification..." — use the same page number as 5.7's own entry
    # since it falls within that section's page range.
    toc_57_idx = find_paragraph_index(
        doc, lambda t: t.startswith("5.7 Model Performance and Evaluation")
    )
    assert toc_57_idx is not None
    toc_57_text = full_text(doc.paragraphs[toc_57_idx])
    page_num = toc_57_text.split("\t")[-1] if "\t" in toc_57_text else ""
    insert_toc_entry_after(
        doc,
        lambda t: t.startswith("5.7 Model Performance and Evaluation"),
        f"5.7.1 Cross-Validation and Robustness Validation\t{page_num}" if page_num else "5.7.1 Cross-Validation and Robustness Validation",
    )
    print(f"[{path.split(chr(92))[-1]}] Added ToC entry for 5.7.1")

    # 1. Insert CV/robustness subsection at the end of Section 5.7
    anchor_idx = find_paragraph_index(
        doc,
        lambda t: t.startswith(
            "ROC-AUC. The tuned pipeline achieves a ROC-AUC of 0.8068 on the held-out test set"
        ),
    )
    assert anchor_idx is not None, f"Could not find Section 5.7 ROC-AUC paragraph in {path}"
    anchor_para = doc.paragraphs[anchor_idx]
    body_style = anchor_para.style

    # find a Heading 3 style reference and its run formatting to copy
    heading3_ref = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 3" and full_text(p).strip():
            heading3_ref = p
            break
    assert heading3_ref is not None

    p3 = insert_paragraph_after(anchor_para, CV_PARA_3, style=body_style)
    p2 = insert_paragraph_after(anchor_para, CV_PARA_2, style=body_style)
    p1 = insert_paragraph_after(anchor_para, CV_PARA_1, style=body_style)
    heading = insert_paragraph_after(anchor_para, CV_HEADING_TEXT, style=heading3_ref.style)
    if heading.runs and heading3_ref.runs:
        heading.runs[0].bold = heading3_ref.runs[0].bold
        heading.runs[0].font.size = heading3_ref.runs[0].font.size
    print(f"[{path.split(chr(92))[-1]}] Inserted CV/robustness subsection after paragraph {anchor_idx}")

    # Re-fetch doc paragraphs are now stale indices for subsequent steps since
    # we inserted elements; but python-docx's Document.paragraphs re-walks the
    # XML tree each time it's accessed, so doc.paragraphs below reflects the
    # current state correctly - no need to reload.

    # 2. Insert permutation-importance paragraph in Section 5.10
    perm_anchor_idx = find_paragraph_index(
        doc,
        lambda t: t.startswith("One result runs against expectation and should be acknowledged"),
    )
    assert perm_anchor_idx is not None, f"Could not find Section 5.10 Overall_Preparedness_Index paragraph in {path}"
    perm_anchor_para = doc.paragraphs[perm_anchor_idx]
    insert_paragraph_after(perm_anchor_para, PERM_IMPORTANCE_PARA, style=perm_anchor_para.style)
    print(f"[{path.split(chr(92))[-1]}] Inserted permutation-importance paragraph after paragraph {perm_anchor_idx}")

    # 3. Fix Section 6.2's two stale limitation paragraphs
    old_464 = (
        "Incomplete algorithmic comparison. Only XGBoost was trained and evaluated. The selection of the "
        "algorithm is justified on established properties in Section 5.8, but no measured comparison against "
        "Logistic Regression, Decision Tree, Random Forest or Support Vector Machine has been performed on this "
        "data, so no claim of comparative superiority is made."
    )
    new_464 = (
        "Incomplete algorithmic comparison. Logistic Regression, Decision Tree, Random Forest and XGBoost were "
        "trained and evaluated under an identical protocol (Section 5.8, Table 5.6); Support Vector Machine was "
        "excluded on computational-cost grounds. The four compared algorithms cluster within two percentage "
        "points of each other on accuracy, with tuned Logistic Regression marginally exceeding tuned XGBoost, so "
        "no claim of XGBoost's outright comparative superiority is made \u2014 its selection instead rests on native "
        "feature-importance output and capacity for further feature engineering, as argued in Section 5.8."
    )
    old_465 = (
        "Uncalibrated precision-recall trade-off. Three imbalance corrections were applied simultaneously and "
        "their combined effect appears to over-correct toward the minority class, yielding 0.36 precision on "
        "flagged graduates. Because no threshold-independent measure was computed, it is not currently possible "
        "to determine how much of the observed performance reflects the model\u2019s discrimination and how much "
        "reflects the threshold choice."
    )
    new_465 = (
        "Precision\u2013recall trade-off. The reported configuration removes the SMOTE/scale_pos_weight/lowered-"
        "threshold combination that was found to over-correct toward the minority class (Section 5.6\u20135.7) and "
        "instead reports minority-class precision of 0.65 at recall of 0.35 under the standard 0.50 threshold, "
        "with ROC-AUC of 0.8068 as the threshold-independent discrimination measure. It remains true that a "
        "precision-recall curve, which is more informative than a single ROC-AUC figure under this degree of "
        "class imbalance, has not been plotted, so a practitioner wishing to choose a different deployment "
        "threshold cannot yet see the full trade-off curve made explicit; this is recorded as outstanding in "
        "Section 5.15 and recommended in Section 6.3."
    )

    for i, p in enumerate(doc.paragraphs):
        t = full_text(p)
        if old_464 in t:
            replace_paragraph_text(p, old_464, new_464)
            print(f"[{path.split(chr(92))[-1]}] Fixed stale 'Incomplete algorithmic comparison' at paragraph {i}")
        elif old_465 in t:
            replace_paragraph_text(p, old_465, new_465)
            print(f"[{path.split(chr(92))[-1]}] Fixed stale 'Uncalibrated precision-recall' at paragraph {i}")

    # 4. Section 5.14: bump "Six limitations" -> "Eight limitations", append two new ones
    fixed_intro = fixed_tail = False
    fixed_fourth = False
    for p in doc.paragraphs:
        t = full_text(p)
        if OLD_LIM_INTRO in t:
            replace_paragraph_text(p, OLD_LIM_INTRO, NEW_LIM_INTRO)
            fixed_intro = True
        if OLD_LIM_FOURTH in t:
            replace_paragraph_text(p, OLD_LIM_FOURTH, NEW_LIM_FOURTH)
            fixed_fourth = True
        if OLD_LIM_TAIL in t:
            replace_paragraph_text(p, OLD_LIM_TAIL, NEW_LIM_TAIL)
            fixed_tail = True
    assert fixed_intro, f"Could not find limitations intro sentence in {path}"
    assert fixed_fourth, f"Could not find limitations 'Fourth' sentence in {path}"
    assert fixed_tail, f"Could not find limitations tail paragraph in {path}"
    print(f"[{path.split(chr(92))[-1]}] Updated Section 5.14 limitations count and added two new limitations")

    doc.save(path)
    print(f"[{path.split(chr(92))[-1]}] Saved.\n")


for path in PATHS:
    process_file(path)

print("Phase 3 findings applied to both documents.")
