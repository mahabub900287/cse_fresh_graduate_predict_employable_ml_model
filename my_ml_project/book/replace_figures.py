"""
Replaces the blob of specific embedded images in Report-02-UPDATED.docx
in place, and resizes each inline shape's displayed width/height to match
the new image's true aspect ratio (avoiding stretched/squashed figures),
while keeping every other part of the document (paragraph position, runs,
captions) untouched.
"""
import docx
from docx.shared import Inches
from PIL import Image

PATH = r"d:\laragon\www\cse_fresh_graduate_predict_employable_ml_model\my_ml_project\book\Report-02-UPDATED.docx"
BOOK_DIR = r"d:\laragon\www\cse_fresh_graduate_predict_employable_ml_model\my_ml_project\book"

doc = docx.Document(PATH)

USABLE_WIDTH_IN = 6.27  # page width minus 1.2" + 0.8" margins, from section geometry
MAX_HEIGHT_IN = 8.5     # keep a tall figure from overflowing a single page badly

# paragraph_index -> (new image file, target display width in inches, or None for auto)
REPLACEMENTS = {
    197: (f"{BOOK_DIR}\\figure_3_1_incremental_model.png", 6.2),
    204: (f"{BOOK_DIR}\\figure_3_2_incremental_system.png", 6.2),
    241: (f"{BOOK_DIR}\\figure_4_2_dfd_level0.png", 5.2),
    244: (f"{BOOK_DIR}\\figure_4_3_dfd_level1.png", 6.2),
    267: (f"{BOOK_DIR}\\figure_5_1_dataset.jpg", 6.2),
    269: (f"{BOOK_DIR}\\figure_5_2_dataset.jpg", 6.2),
    350: (f"{BOOK_DIR}\\webview_screenshot_form_final.png", 4.6),
    353: (f"{BOOK_DIR}\\webview_screenshot_result_final.png", 4.6),
}


def full_text(p):
    return "".join(r.text for r in p.runs)


for idx, (new_path, target_w_in) in REPLACEMENTS.items():
    p = doc.paragraphs[idx]
    caption = full_text(doc.paragraphs[idx + 1]) if idx + 1 < len(doc.paragraphs) else "(no caption found)"
    blips = p._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
    assert blips, f"No image found at paragraph {idx} (expected caption: {caption})"
    rId = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
    image_part = p.part.related_parts[rId]

    with open(new_path, "rb") as f:
        new_blob = f.read()
    image_part._blob = new_blob

    img = Image.open(new_path)
    aspect = img.height / img.width
    disp_w = min(target_w_in, USABLE_WIDTH_IN)
    disp_h = disp_w * aspect
    if disp_h > MAX_HEIGHT_IN:
        disp_h = MAX_HEIGHT_IN
        disp_w = disp_h / aspect

    # Find the drawing's extent elements (both wp:extent and a:ext under xfrm) and update them.
    drawing = p._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline")
    assert drawing, f"No inline drawing wrapper found at paragraph {idx}"
    inline_el = drawing[0]
    cx = Inches(disp_w).emu if hasattr(Inches(disp_w), "emu") else int(disp_w * 914400)
    cy = int(disp_h * 914400)

    ns_wp = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    ns_a = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    extent = inline_el.find(f"{ns_wp}extent")
    extent.set("cx", str(cx))
    extent.set("cy", str(cy))
    for ext in inline_el.findall(f".//{ns_a}ext"):
        ext.set("cx", str(cx))
        ext.set("cy", str(cy))

    print(f"Paragraph {idx} ({caption[:55]!r}): {new_path.split(chr(92))[-1]} -> {disp_w:.2f}in x {disp_h:.2f}in")

doc.save(PATH)
print("\nSaved.")
