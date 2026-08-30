"""
Inserts a new subsection into Section 5.9 (System Implementation and Output)
describing the local browser-based web view (webview.html) built for this
project, with two embedded screenshots, and corrects the surrounding prose
that previously stated no front end had been built.
"""
import copy

import docx
from docx.shared import Inches

PATH = r"d:\laragon\www\cse_fresh_graduate_predict_employable_ml_model\my_ml_project\book\Report-02-UPDATED.docx"
FORM_SCREENSHOT = r"d:\laragon\www\cse_fresh_graduate_predict_employable_ml_model\my_ml_project\book\webview_screenshot_form.png"
RESULT_SCREENSHOT = r"d:\laragon\www\cse_fresh_graduate_predict_employable_ml_model\my_ml_project\book\webview_screenshot_result.png"

doc = docx.Document(PATH)


def full_text(p):
    return "".join(r.text for r in p.runs)


def replace_in_paragraph(paragraph, old, new):
    full = full_text(paragraph)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


# ---------------------------------------------------------------------------
# 1. Correct the "no web-based front end has been built" claim
# ---------------------------------------------------------------------------
old_342 = (
    "The selected model was made operational as a documented HTTP API so that its predictions could be consumed directly. "
    "At the current stage of the project no web-based front end has been built: a student\u2019s data is submitted directly "
    "to the trained model through the API, and the resulting employability evaluation is obtained in the same request, "
    "without any account, registration or login."
)
new_342 = (
    "The selected model was made operational as a documented HTTP API so that its predictions could be consumed directly. "
    "In addition to the API itself, a local single-page web view (webview.html) was built so that a graduate can submit "
    "a profile and read the resulting evaluation without constructing an HTTP request by hand; it is described later in "
    "this section. A student\u2019s data is submitted to the trained model through the API \u2014 either directly, or via this "
    "web view \u2014 and the resulting employability evaluation is obtained in the same request, without any account, "
    "registration or login."
)
found = False
for p in doc.paragraphs:
    if replace_in_paragraph(p, old_342, new_342):
        found = True
        break
print("Paragraph 342 (no-frontend claim) updated:", found)

# ---------------------------------------------------------------------------
# 2. Locate the paragraph just before "5.10 Model Interpretation..." (the
#    FR6/FR7 paragraph) and insert the new web-view subsection immediately
#    after it, before the section heading.
# ---------------------------------------------------------------------------
anchor_text = "5.10 Model Interpretation: Feature Importance and Feature Priority"
anchor_idx = None
for i, p in enumerate(doc.paragraphs):
    if full_text(p).strip() == anchor_text:
        anchor_idx = i
        break
assert anchor_idx is not None, "Could not find Section 5.10 heading"
anchor_paragraph = doc.paragraphs[anchor_idx]
print("Found Section 5.10 heading at paragraph index", anchor_idx)

# Grab a heading-styled paragraph to copy formatting from (find "5.9 ..." heading)
heading_59 = None
for p in doc.paragraphs:
    if full_text(p).strip().startswith("5.9 System Implementation"):
        heading_59 = p
        break
assert heading_59 is not None

heading_style = heading_59.style
body_style = doc.paragraphs[anchor_idx - 1].style  # a normal body paragraph style
caption_style = doc.paragraphs[345].style  # "Figure 5.11: ..." caption style


def insert_paragraph_before(reference_paragraph, text="", style=None):
    new_p = docx.oxml.ns.qn  # placeholder to keep import usage explicit
    new_element = copy.deepcopy(reference_paragraph._p)
    # Clear runs from the copy, then reuse it as a template inserted before reference
    for child in list(new_element):
        if child.tag.endswith("}r"):
            new_element.remove(child)
    reference_paragraph._p.addprevious(new_element)
    from docx.text.paragraph import Paragraph
    new_paragraph = Paragraph(new_element, reference_paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


# Insert (in order) before the "5.10 ..." heading paragraph:
sub_heading = insert_paragraph_before(
    anchor_paragraph,
    "Local Web View for Interactive Predictions",
    style=None,
)
# Give it the same bold/emphasis as a minor heading by matching heading_59's run formatting where possible
if heading_59.runs:
    src_run = heading_59.runs[0]
    if sub_heading.runs:
        dst_run = sub_heading.runs[0]
        dst_run.bold = src_run.bold
        dst_run.font.size = src_run.font.size

body_para_1 = insert_paragraph_before(
    anchor_paragraph,
    (
        "To make the API usable without a separate HTTP client, a self-contained local web page (webview.html) was built "
        "alongside the FastAPI service. The page is opened directly in a browser from the local file system and calls "
        "the same /predict endpoint documented above \u2014 the FastAPI application was extended with permissive CORS "
        "middleware specifically so that a page opened via the file:// protocol can call it. The form on the left "
        "collects the same profile attributes used for prediction, grouped into Academic, Technical & Portfolio, "
        "Experience & Readiness, and Soft Skills sections, plus the six demographic and contextual fields that are "
        "accepted by the request schema but excluded from the model itself (Section 5.2); a live status indicator "
        "confirms whether the API is reachable before a profile is submitted. Figure 5.11a shows the form in its "
        "initial state."
    ),
    style=body_style,
)

image_para_1 = insert_paragraph_before(anchor_paragraph, style=None)
run_img1 = image_para_1.add_run()
run_img1.add_picture(FORM_SCREENSHOT, width=Inches(5.8))
image_para_1.alignment = doc.paragraphs[344].alignment

caption_1 = insert_paragraph_before(
    anchor_paragraph,
    "Figure 5.11a: The local web view's profile-submission form, showing its idle result panel before a check is run",
    style=caption_style,
)

body_para_2 = insert_paragraph_before(
    anchor_paragraph,
    (
        "On submission, the page sends the profile to /predict and renders the response on the right: a verdict badge "
        "(Employable or Not Employable), a probability gauge, and, for a Not Employable verdict, the same rule-based "
        "gap-analysis list produced by the API (Section 5.9). Figure 5.11b shows this state for a deliberately weak "
        "profile \u2014 low CGPA, no internships, a below-median interview score \u2014 which the trained pipeline scores at "
        "14 per cent employability probability and for which it lists eight concrete gaps. This satisfies FR5 (Table "
        "4.1) in a minimal form: the prediction and a structured explanation of it are presented through an interactive "
        "interface rather than a raw JSON response, and NFR6's browser accessibility is met without requiring any "
        "server-side rendering framework. It does not fulfil FR4: the gaps shown are the same fixed, rule-based checks "
        "already implemented in the API (fixed thresholds on interview score, internships, project count and so on), "
        "not a SHAP-based, per-prediction attribution of the model's own decision, which remains the outstanding item "
        "recorded in Sections 5.9, 5.15 and 6.3. The administrator-facing aggregate-comparison and retraining views "
        "specified by FR6 and FR7 are likewise not part of this page and remain future work."
    ),
    style=body_style,
)

image_para_2 = insert_paragraph_before(anchor_paragraph, style=None)
run_img2 = image_para_2.add_run()
run_img2.add_picture(RESULT_SCREENSHOT, width=Inches(5.8))
image_para_2.alignment = doc.paragraphs[344].alignment

caption_2 = insert_paragraph_before(
    anchor_paragraph,
    "Figure 5.11b: The local web view after a submission, showing the verdict badge, probability gauge and gap-analysis list for a weak profile",
    style=caption_style,
)

blank_after = insert_paragraph_before(anchor_paragraph, "", style=body_style)

print("Inserted web-view subsection with two figures before Section 5.10.")

# ---------------------------------------------------------------------------
# 3. Soften the FR5/NFR6 "not yet exposed through any user-facing interface"
#    framing in the FR6/FR7 paragraph, since FR5 is now partially met.
# ---------------------------------------------------------------------------
old_347 = (
    "The rendering of attribute contributions as visual charts for non-technical users (FR5, NFR6) likewise awaits the dashboard described there."
)
new_347 = (
    "The rendering of attribute contributions as visual charts for non-technical users (FR5, NFR6) is now partially "
    "met by the local web view described above, which renders the verdict and gap list visually; a full SHAP-based, "
    "per-prediction attribution chart still awaits the dashboard described there."
)
found2 = False
for p in doc.paragraphs:
    if replace_in_paragraph(p, old_347, new_347):
        found2 = True
        break
print("Paragraph 347 (FR5/NFR6 framing) updated:", found2)

doc.save(PATH)
print("Saved.")
