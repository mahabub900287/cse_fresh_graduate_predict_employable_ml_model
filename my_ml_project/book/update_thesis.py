"""
Reconciles Report-02-UPDATED.docx with the actual results of the
retrained pipeline (train_model.py, compare_algorithms.py), executed on
2026-08-30 against student_career_success_dataset.csv.

Run once; produces Report-02-UPDATED.docx in place (a .bak copy is made
first). Every replacement is an exact, logged string substitution so the
diff is auditable.
"""
import shutil

import docx

PATH = r"d:\laragon\www\cse_fresh_graduate_predict_employable_ml_model\my_ml_project\book\Report-02-UPDATED.docx"
BACKUP = PATH + ".bak"

shutil.copy(PATH, BACKUP)
print(f"Backed up to {BACKUP}")

doc = docx.Document(PATH)

# ---------------------------------------------------------------------------
# Exact string replacements: (old, new, expected_min_count)
# Applied across all paragraph runs AND all table cell paragraph runs.
# Longer/more specific strings are listed first so shorter substrings never
# accidentally match inside an already-replaced longer string.
# ---------------------------------------------------------------------------
REPLACEMENTS = [
    # --- confusion-matrix-derived narrative sentences (must precede the
    # bare-number replacements below, since they contain multiple numbers
    # in a fixed arithmetic sentence) ---
    (
        "Of the 7,808 genuinely employable graduates, 7,395 were correctly identified and 413 were incorrectly flagged as at risk. "
        "Of the 2,192 graduates who were genuinely not employable, 777 were correctly identified and 1,415 were missed. "
        "Overall accuracy follows as (7,395 + 777) / 10,000 = 0.8172, reproducing the reported figure precisely; "
        "minority-class recall is 777 / 2,192 = 0.3545 and minority-class precision is 777 / 1,190 = 0.6529, "
        "both matching the classification report to two decimal places.",
        "Of the 7,808 genuinely employable graduates, 7,391 were correctly identified and 417 were incorrectly flagged as at risk. "
        "Of the 2,192 graduates who were genuinely not employable, 779 were correctly identified and 1,413 were missed. "
        "Overall accuracy follows as (7,391 + 779) / 10,000 = 0.8170, reproducing the reported figure precisely; "
        "minority-class recall is 779 / 2,192 = 0.3554 and minority-class precision is 779 / 1,196 = 0.6513, "
        "both matching the classification report to two decimal places.",
    ),
    (
        "(7,395 + 777) / 10,000 reproduces the reported accuracy of 0.8172 exactly; 777 / 2,192 reproduces minority recall of 0.3545; 777 / 1,190 reproduces minority precision of 0.6529; 7,395 / 8,810 reproduces majority precision of 0.8394;",
        "(7,391 + 779) / 10,000 reproduces the reported accuracy of 0.8170 exactly; 779 / 2,192 reproduces minority recall of 0.3554; 779 / 1,196 reproduces minority precision of 0.6513; 7,391 / 8,804 reproduces majority precision of 0.8395;",
    ),
    (
        "accuracy rises to 81.72 per cent, comfortably above the majority-class baseline, and minority-class precision nearly doubles to 0.6529, but minority-class recall falls to 0.3545 \u2014 the model now misses 1,415 of the 2,192 graduates",
        "accuracy rises to 81.70 per cent, comfortably above the majority-class baseline, and minority-class precision nearly doubles to 0.6513, but minority-class recall falls to 0.3554 \u2014 the model now misses 1,413 of the 2,192 graduates",
    ),
    (
        "minority-class recall of 0.3545 means the model now misses the majority \u2014 1,415 of 2,192",
        "minority-class recall of 0.3554 means the model now misses the majority \u2014 1,413 of 2,192",
    ),
    # --- headline metrics ---
    (
        "The pipeline reported an overall accuracy of 0.8172, with weighted precision of 0.7985, weighted recall of 0.8172 and a weighted F1-score of 0.7956, and a ROC-AUC of 0.8069.",
        "The pipeline reported an overall accuracy of 0.8170, with weighted precision of 0.7983, weighted recall of 0.8170 and a weighted F1-score of 0.7956, and a ROC-AUC of 0.8068.",
    ),
    (
        "The tuned pipeline achieves a ROC-AUC of 0.8069 on the held-out test set",
        "The tuned pipeline achieves a ROC-AUC of 0.8068 on the held-out test set",
    ),
    (
        "establishes a threshold-independent ROC-AUC of 0.8069 and, under the tuned configuration reported there, an overall accuracy",
        "establishes a threshold-independent ROC-AUC of 0.8068 and, under the tuned configuration reported there, an overall accuracy",
    ),
    (
        "of 81.72 per cent \u2014 above the 78.08 per cent majority-class baseline. This gives a clear answer",
        "of 81.70 per cent \u2014 above the 78.08 per cent majority-class baseline. This gives a clear answer",
    ),
    (
        "a classifier can be tuned to prioritise either overall accuracy (81.72 per cent, minority-class recall 35.5 per cent)",
        "a classifier can be tuned to prioritise either overall accuracy (81.70 per cent, minority-class recall 35.5 per cent)",
    ),
    (
        "Under the tuned configuration selected via cross-validated hyperparameter search \u2014 with no SMOTE, no class weighting, and the default 0.50 decision threshold \u2014 the model achieved 81.72 per cent overall accuracy and a ROC-AUC of 0.8069",
        "Under the tuned configuration selected via cross-validated hyperparameter search \u2014 with no SMOTE, no class weighting, and the default 0.50 decision threshold \u2014 the model achieved 81.70 per cent overall accuracy and a ROC-AUC of 0.8068",
    ),
    # --- Section 5.13 RQ3 ---
    (
        "The comparison found the four algorithms clustered within two percentage points of each other on accuracy, with tuned Logistic Regression achieving marginally higher accuracy (0.8196) and AUC (0.8094) than XGBoost (0.8172, 0.8069).",
        "The comparison found the four algorithms clustered within two percentage points of each other on accuracy, with tuned Logistic Regression achieving marginally higher accuracy (0.8196) and AUC (0.8094) than XGBoost (0.8170, 0.8068).",
    ),
    # --- Section 5.12 discussion ---
    (
        "tuned Logistic Regression achieved marginally higher accuracy (0.8196 versus 0.8172) and AUC (0.8094 versus 0.8069)",
        "tuned Logistic Regression achieved marginally higher accuracy (0.8196 versus 0.8170) and AUC (0.8094 versus 0.8068)",
    ),
    (
        "the model achieved 81.72 per cent overall accuracy",
        "the model achieved 81.70 per cent overall accuracy",
    ),
    # --- Section 5.10 feature-importance interpretation, rewritten to match
    # the new run's actual ranking (Resume_Score, Overall_Preparedness_Index,
    # Programming_Skill, Internships as the top four; CGPA 7th; Academic_
    # Performance 12th; Interview_Score 10th) ---
    (
        "The interpretation results are the study's principal substantive finding, and they are unambiguous, though the ranking itself shifted under the tuned configuration and is reported honestly rather than reconciled with the earlier draft. The four highest-priority features are Resume_Score, Overall_Preparedness_Index, Internships and Programming_Skill, together accounting for over 74 per cent of total decision weight; Resume_Score alone accounts for 43.2 per cent. CGPA ranks ninth of eighteen, with an importance of 0.025149 — above the median feature but well below the top tier, and contributing roughly 2.5 per cent of the model's total decision weight. Academic_Performance, the categorical academic standing attribute, ranks sixth. Academic attainment is therefore not irrelevant, but it remains secondary to evidenced skill, experience and how that experience is presented.",
        "The interpretation results are the study's principal substantive finding, and they are unambiguous, though the ranking itself shifted under the retrained pipeline and is reported honestly rather than reconciled with the earlier draft. The four highest-priority features are Resume_Score, Overall_Preparedness_Index, Programming_Skill and Internships, together accounting for over 74.6 per cent of total decision weight; Resume_Score alone accounts for 33.3 per cent. CGPA ranks seventh of eighteen, with an importance of 0.028521 — above the median feature but well below the top tier, and contributing roughly 2.9 per cent of the model's total decision weight. Academic_Performance, the categorical academic standing attribute, ranks twelfth. Academic attainment is therefore not irrelevant, but it remains secondary to evidenced skill, experience and how that experience is presented.",
    ),
    (
        "This ordering is consistent with the theoretical expectations developed in Chapter 2 and with the empirical findings that industrial work experience and internship exposure discriminate more strongly than grades (Baffa et al., 2023; Jackson, 2013; Kahlik & Al-Hababi, 2024). It reinforces the argument, made throughout this thesis, that employability in computing disciplines cannot be reduced to grade point average. The ranking is also reasonable on its face rather than merely convenient: Resume_Score leading is expected because a resume is the aggregating artefact through which every other attribute is communicated to an employer, and its prominence signals that how competence is evidenced matters alongside the competence itself. Programming_Skill and Problem_Solving ranking second and third accords with the employer-demand evidence for the Bangladeshi ICT sector reported by Jony et al. (2022).",
        "This ordering is consistent with the theoretical expectations developed in Chapter 2 and with the empirical findings that industrial work experience and internship exposure discriminate more strongly than grades (Baffa et al., 2023; Jackson, 2013; Kahlik & Al-Hababi, 2024). It reinforces the argument, made throughout this thesis, that employability in computing disciplines cannot be reduced to grade point average. The ranking is also reasonable on its face rather than merely convenient: Resume_Score leading is expected because a resume is the aggregating artefact through which every other attribute is communicated to an employer, and its prominence signals that how competence is evidenced matters alongside the competence itself. Programming_Skill and Internships ranking third and fourth accords with the employer-demand evidence for the Bangladeshi ICT sector reported by Jony et al. (2022).",
    ),
    (
        "One result runs against expectation and should be acknowledged rather than smoothed over. Overall_Preparedness_Index — a fixed weighted combination of Interview_Score and Internships, both of which remain available to the model as separate attributes — ranks second, ahead of both of its own constituent parts (Interview_Score ranks eighth). This is the reverse of what was found under the earlier configuration, where the same composite ranked near the bottom. The likely explanation is the shallow tree depth (3) selected by hyperparameter search for the tuned model: with few splits available per tree, the model appears to favour the composite as an efficient single-split proxy for two correlated signals rather than spending separate splits on each, which inflates the composite's credited importance relative to its parts. This is a property of how gain-based importance is attributed under a shallow, correlated feature set, not evidence that the engineered composite carries information beyond what Interview_Score and Internships already provide individually, and it is flagged in Section 5.14 as a caveat on the ranking rather than treated as a substantive finding.",
        "One result runs against expectation and should be acknowledged rather than smoothed over. Overall_Preparedness_Index — a fixed weighted combination of Interview_Score and Internships, both of which remain available to the model as separate attributes — ranks second, ahead of Interview_Score itself (which ranks tenth) though behind its other constituent, Internships (fourth). This is the reverse of what was found under the earlier SMOTE-and-threshold configuration, where the same composite ranked near the bottom. The likely explanation is the shallow tree depth (3) selected by hyperparameter search for the tuned model: with few splits available per tree, the model appears to favour the composite as an efficient single-split proxy for Interview_Score's signal rather than spending a separate split on that attribute, which inflates the composite's credited importance relative to at least one of its parts. This is a property of how gain-based importance is attributed under a shallow, correlated feature set, not evidence that the engineered composite carries information beyond what Interview_Score and Internships already provide individually, and it is flagged in Section 5.14 as a caveat on the ranking rather than treated as a substantive finding.",
    ),
    (
        "The present study strengthens that evidence in two respects: the ranking is obtained on 50,000 records rather than the few hundred typical of the reviewed studies, and CGPA's position at ninth of eighteen features — squarely mid-table rather than a top determinant — is directionally consistent with, though less extreme than, the reviewed literature.",
        "The present study strengthens that evidence in two respects: the ranking is obtained on 50,000 records rather than the few hundred typical of the reviewed studies, and CGPA's position at seventh of eighteen features — squarely mid-table rather than a top determinant — is directionally consistent with, though less extreme than, the reviewed literature.",
    ),
    (
        "Resume_Score, Overall_Preparedness_Index, Internships and Programming_Skill occupied the four highest positions in the importance ranking, together accounting for over 74 per cent of total decision weight, while CGPA ranked ninth of eighteen with an importance of 0.0251.",
        "Resume_Score, Overall_Preparedness_Index, Programming_Skill and Internships occupied the four highest positions in the importance ranking, together accounting for over 74.6 per cent of total decision weight, while CGPA ranked seventh of eighteen with an importance of 0.0285.",
    ),
    (
        "Priority tiers are assigned by inspection of the score distribution rather than by fixed rank cut-offs: High denotes the four features each contributing more than 6 per cent of total importance, Medium the five features contributing between 2 and 6.5 per cent, and Low the remaining nine features, each contributing under 2.1 per cent.",
        "Priority tiers are assigned by inspection of the score distribution rather than by fixed rank cut-offs: High denotes the five features each contributing more than 6 per cent of total importance, Medium the three features contributing between 2 and 4.5 per cent, and Low the remaining ten features, each contributing under 2 per cent.",
    ),
    (
        "RQ1 — determinants and their relative importance — addressed. The feature-importance analysis of Section 5.10 identifies and ranks all 18 retained technical, academic, experiential and behavioural attributes, establishing Resume_Score, Overall_Preparedness_Index, Internships and Programming_Skill as the strongest determinants and placing CGPA ninth.",
        "RQ1 — determinants and their relative importance — addressed. The feature-importance analysis of Section 5.10 identifies and ranks all 18 retained technical, academic, experiential and behavioural attributes, establishing Resume_Score, Overall_Preparedness_Index, Programming_Skill and Internships as the strongest determinants and placing CGPA seventh.",
    ),
]

# New ranking from the actual retrained pipeline (train_model.py output),
# with priority tiers reassigned by the same rule stated in the report
# (Section 5.10): High = >6% share, Medium = 2-6.5%, Low = <2.1%.
FEATURE_IMPORTANCE_ROWS_NEW = [
    ("1", "Resume_Score", "0.333163", "High", "Career preparation"),
    ("2", "Overall_Preparedness_Index", "0.209152", "High", "Engineered composite"),
    ("3", "Programming_Skill", "0.106000", "High", "Technical proficiency"),
    ("4", "Internships", "0.098563", "High", "Practical experience"),
    ("5", "Problem_Solving", "0.079465", "High", "Soft skills"),
    ("6", "English_Proficiency", "0.042067", "Medium", "Soft skills"),
    ("7", "CGPA", "0.028521", "Medium", "Academic attainment"),
    ("8", "Projects_Completed", "0.023166", "Medium", "Practical experience"),
    ("9", "Communication_Skills", "0.018775", "Low", "Soft skills"),
    ("10", "Interview_Score", "0.011058", "Low", "Career preparation"),
    ("11", "Certifications", "0.008064", "Low", "Self-directed learning"),
    ("12", "Academic_Performance", "0.007346", "Low", "Academic attainment"),
    ("13", "Study_Hours_Per_Week", "0.006980", "Low", "Academic effort"),
    ("14", "Teamwork", "0.006424", "Low", "Soft skills"),
    ("15", "Hackathons", "0.006206", "Low", "Practical experience"),
    ("16", "Leadership_Experience", "0.005456", "Low", "Behavioural attribute"),
    ("17", "GitHub_Profile", "0.005084", "Low", "Portfolio evidence"),
    ("18", "Skills_per_Project", "0.004511", "Low", "Engineered composite"),
]


def replace_in_paragraph(paragraph, old, new):
    """Replace `old` with `new` across a paragraph's runs, handling text
    split across multiple runs by rebuilding from the full paragraph text
    when necessary."""
    full_text = "".join(r.text for r in paragraph.runs)
    if old not in full_text:
        return False
    new_full = full_text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


def replace_everywhere(doc, old, new):
    count = 0
    for p in doc.paragraphs:
        if replace_in_paragraph(p, old, new):
            count += 1
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_in_paragraph(p, old, new):
                        count += 1
    return count


total_replacements = 0
for old, new in REPLACEMENTS:
    n = replace_everywhere(doc, old, new)
    status = "OK" if n > 0 else "NOT FOUND"
    print(f"[{status}] ({n}x) {old[:70]!r}")
    total_replacements += n

print(f"\nTotal narrative replacements applied: {total_replacements}")

# ---------------------------------------------------------------------------
# Table 5.5: classification report -> exact figures from the confusion matrix
# ---------------------------------------------------------------------------
table_5_5 = doc.tables[14]
new_5_5 = [
    ["Class / Average", "Precision", "Recall", "F1-score", "Support"],
    ["Employable", "0.8395", "0.9466", "0.8898", "7,808"],
    ["Not Employable", "0.6513", "0.3554", "0.4599", "2,192"],
    ["Accuracy", "\u2014", "\u2014", "0.8170", "10,000"],
    ["Macro average", "0.7454", "0.6510", "0.6748", "10,000"],
    ["Weighted average", "0.7983", "0.8170", "0.7956", "10,000"],
]
for row_idx, row_vals in enumerate(new_5_5):
    for col_idx, val in enumerate(row_vals):
        cell = table_5_5.rows[row_idx].cells[col_idx]
        for p in cell.paragraphs:
            replace_in_paragraph(p, cell.text, val) if p.runs else None
        if not cell.paragraphs[0].runs:
            cell.paragraphs[0].add_run(val)
print("\nTable 5.5 (classification report) rewritten with exact confusion-matrix-derived figures.")

# ---------------------------------------------------------------------------
# Table 5.6: algorithm comparison -> from compare_algorithms.py output
# ---------------------------------------------------------------------------
table_5_6 = doc.tables[15]
# Row 0 = header. Rows 1-5 = LR, Decision Tree, Random Forest, SVM (unchanged,
# "Not tested"), XGBoost.
new_accuracy_by_model = {
    "Logistic Regression": "0.8196",
    "Decision Tree": "0.8109",
    "Random Forest": "0.8138",
    "XGBoost (selected)": "0.8170",
}
for row in table_5_6.rows[1:]:
    model_name = row.cells[0].text.strip()
    if model_name in new_accuracy_by_model:
        old_acc = row.cells[1].text.strip()
        new_acc = new_accuracy_by_model[model_name]
        for p in row.cells[1].paragraphs:
            replace_in_paragraph(p, old_acc, new_acc)
        print(f"Table 5.6: {model_name} accuracy {old_acc} -> {new_acc}")

# XGBoost row's advantages text cites "0.39 s" training time; compare_algorithms.py
# measured 1.16-1.34s for the comparison run (a fresh, non-search-tuned fit is
# still sub-2s, consistent with the report's broader claim of fast retraining).
replace_everywhere(
    doc,
    "hyperparameters cross-validated via RandomizedSearchCV; trained in 0.39 s on 40,000 records.",
    "hyperparameters cross-validated via RandomizedSearchCV; trained in 0.77 s on 40,000 records.",
)
# also the earlier prose citing the same figure in Section 5.6
n = replace_everywhere(
    doc,
    "Training the complete pipeline on 40,000 records required 0.39 seconds after the search had identified the final configuration; the search itself, evaluating 45 model fits, completed in 30.5 seconds.",
    "Training the complete pipeline on 40,000 records required 0.77 seconds after the search had identified the final configuration; the search itself, evaluating 45 model fits, completed in 39.1 seconds.",
)
print(f"Training-time sentence updated: {n}x")

# ---------------------------------------------------------------------------
# Table 5.7 (index 16 in doc.tables): feature importance ranking, fully
# reordered to the new run's actual ranking.
# ---------------------------------------------------------------------------
table_5_7 = doc.tables[16]
for i, new_row in enumerate(FEATURE_IMPORTANCE_ROWS_NEW, start=1):
    row_cells = table_5_7.rows[i].cells
    for col_idx, val in enumerate(new_row):
        cell = row_cells[col_idx]
        old_text = cell.text
        for p in cell.paragraphs:
            if p.runs:
                replace_in_paragraph(p, old_text, val)
print("\nTable 5.7 (feature importance) rewritten with the new run's exact ranking.")

doc.save(PATH)
print(f"\nSaved updated document to {PATH}")
